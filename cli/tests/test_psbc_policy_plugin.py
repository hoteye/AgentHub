import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch
import zipfile
import builtins

from docx import Document

ROOT = Path(__file__).resolve().parents[2]

from cli.agent_cli.host.plugin_manager import PluginManager
from cli.agent_cli.host_platform import current_host_platform
from cli.agent_cli.models import ToolEvent
from cli.agent_cli.provider import _command_for_tool_call, _tool_specs
from cli.agent_cli.providers.config_catalog import ProviderConfig
from cli.agent_cli.providers.tool_specs import responses_provider_tool_specs
from cli.agent_cli.runtime import AgentCliRuntime
from cli.agent_cli.tools import ToolRegistry
from cli.tests.provider_boundary_test_support import provider_status_path_fields

class _PluginFakeAgent:
    def provider_status(self):
        return {
            "provider_ready": "true",
            "provider_name": "deepseek",
            "provider_planner": "deepseek_chat",
            "provider_model": "test-model",
            "provider_tools": "tool-calls",
            "provider_label": "deepseek | test-model | tool-calls",
            "provider_base_url": "http://test",
            "provider_source": "test",
            **provider_status_path_fields(),
        }

    def plan(self, text, history=None, *, tool_executor=None, attachments=None):
        raise AssertionError("LLM planner should not be used in psbc_policy plugin tests")

class PsbcPolicyPluginTest(unittest.TestCase):
    @staticmethod
    def _psbc_plugin_path() -> Path:
        return ROOT / "plugins" / "psbc_policy"

    def _build_plugin_zip(self, root: Path) -> Path:
        source_dir = self._psbc_plugin_path()
        zip_path = root / "psbc_policy.zip"
        with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
            for path in source_dir.rglob("*"):
                if not path.is_file():
                    continue
                if "__pycache__" in path.parts or path.suffix in {".pyc", ".pyo"}:
                    continue
                zf.write(path, arcname=str(path.relative_to(source_dir.parent)))
        return zip_path

    @staticmethod
    def _openai_model_tool_names(manager: PluginManager) -> list[str]:
        specs = responses_provider_tool_specs(
            ProviderConfig(
                model="gpt-5.4",
                api_key="test",
                interaction_profile="codex_openai",
                planner_kind="openai_responses",
                wire_api="responses",
            ),
            current_host_platform(),
            plugin_manager_factory=lambda: manager,
        )
        return [str((item.get("function") or {}).get("name") or item.get("name") or "") for item in specs]

    def test_plugin_manager_discovers_psbc_policy(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            manager = PluginManager(state_path=Path(tmpdir) / "plugin_state.json")
            plugins = manager.list_plugins()
            self.assertTrue(
                any(
                    item["name"] == "psbc_policy"
                    and not item["enabled"]
                    and item["commercial"]
                    and item["distribution"] == "external"
                    and item["plugin_kind"] == "business"
                    for item in plugins
                )
            )
            self.assertFalse(any(item["name"] == "policy_doc_search" for item in manager.command_specs()))
            self.assertNotIn("policy_query", self._openai_model_tool_names(manager))

            enabled = manager.enable_plugin("psbc_policy")
            self.assertTrue(enabled["ok"])
            self.assertTrue(any(item["name"] == "policy_doc_search" for item in manager.command_specs()))
            self.assertTrue(any(item["name"] == "policy_query" for item in manager.command_specs()))
            provider_tool_names = [item["function"]["name"] for item in manager.provider_tool_specs()]
            self.assertEqual(provider_tool_names, ["policy_query"])
            openai_tool_names = self._openai_model_tool_names(manager)
            self.assertIn("policy_query", openai_tool_names)
            self.assertNotIn("policy_doc_import", openai_tool_names)
            self.assertNotIn("policy_doc_list", openai_tool_names)
            self.assertNotIn("policy_doc_search", openai_tool_names)
            self.assertNotIn("policy_doc_read", openai_tool_names)

    def test_provider_maps_policy_tool_call_to_plugin_command(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            manager = PluginManager(state_path=Path(tmpdir) / "plugin_state.json")
            manager.enable_plugin("psbc_policy")
            with patch("cli.agent_cli.providers.tool_calls.PluginManager", return_value=manager):
                command = _command_for_tool_call(
                    "policy_doc_search",
                    {"query": "长期不使用 运维用户", "limit": 5},
                    current_host_platform(),
                )
        self.assertIsNotNone(command)
        self.assertIn("/policy_doc_search", command or "")
        self.assertIn("--query", command or "")
        self.assertIn("长期不使用 运维用户", command or "")
        self.assertIn("--limit 5", command or "")

    def test_provider_tool_specs_expose_only_unified_policy_query_to_llm(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            manager = PluginManager(state_path=Path(tmpdir) / "plugin_state.json")
            manager.enable_plugin("psbc_policy")
            specs = responses_provider_tool_specs(
                ProviderConfig(
                    model="gpt-5.4",
                    api_key="test",
                    interaction_profile="codex_openai",
                    planner_kind="openai_responses",
                    wire_api="responses",
                ),
                current_host_platform(),
                plugin_manager_factory=lambda: manager,
            )
        tool_names = [str((item.get("function") or {}).get("name") or item.get("name") or "") for item in specs]
        self.assertIn("list_dir", tool_names)
        self.assertIn("grep_files", tool_names)
        self.assertIn("read_file", tool_names)
        self.assertIn("policy_query", tool_names)
        self.assertNotIn("policy_doc_import", tool_names)
        self.assertNotIn("policy_doc_list", tool_names)
        self.assertNotIn("policy_doc_search", tool_names)
        self.assertNotIn("policy_doc_read", tool_names)

    def test_psbc_policy_injects_and_cleans_up_slash_and_llm_tools(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            manager = PluginManager(state_path=Path(tmpdir) / "plugin_state.json")

            self.assertFalse(any(item["name"] == "policy_doc_search" for item in manager.command_specs()))
            self.assertFalse(any(item["name"] == "policy_query" for item in manager.command_specs()))
            self.assertNotIn("policy_query", self._openai_model_tool_names(manager))

            enabled = manager.enable_plugin("psbc_policy")
            self.assertTrue(enabled["ok"])
            command_names = {item["name"] for item in manager.command_specs()}
            tool_names = {item["name"] for item in manager.tool_specs()}
            self.assertTrue({"policy_doc_import", "policy_doc_list", "policy_doc_search", "policy_doc_read", "policy_query"}.issubset(command_names))
            self.assertTrue({"policy_doc_import", "policy_doc_list", "policy_doc_search", "policy_doc_read", "policy_query"}.issubset(tool_names))
            openai_tool_names = self._openai_model_tool_names(manager)
            self.assertIn("policy_query", openai_tool_names)
            self.assertNotIn("policy_doc_search", openai_tool_names)

            disabled = manager.disable_plugin("psbc_policy")
            self.assertTrue(disabled["ok"])
            command_names = {item["name"] for item in manager.command_specs()}
            tool_names = {item["name"] for item in manager.tool_specs()}
            self.assertFalse({"policy_doc_import", "policy_doc_list", "policy_doc_search", "policy_doc_read", "policy_query"} & command_names)
            self.assertFalse({"policy_doc_import", "policy_doc_list", "policy_doc_search", "policy_doc_read", "policy_query"} & tool_names)
            self.assertNotIn("policy_query", self._openai_model_tool_names(manager))

    def test_runtime_executes_policy_search_via_plugin_command_path(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            tools = ToolRegistry()
            tools._plugin_manager = PluginManager(state_path=Path(tmpdir) / "plugin_state.json")
            tools._plugin_manager.enable_plugin("psbc_policy")
            runtime = AgentCliRuntime(tools=tools, agent=_PluginFakeAgent())

            def _unexpected_host_search(*args, **kwargs):
                raise AssertionError("host policy_doc_search branch should not execute")

            def _fake_plugin_search(**kwargs):
                return ToolEvent(
                    name="policy_doc_search",
                    ok=True,
                    summary="policy search ok",
                    payload={"query": kwargs.get("query"), "match_count": 1},
                )

            tools.policy_doc_search = _unexpected_host_search
            tools._plugin_manager._tools["policy_doc_search"].handler = _fake_plugin_search

            response = runtime.handle_prompt("/policy_doc_search query 长期不使用 limit 3")
            self.assertEqual([event.name for event in response.tool_events], ["policy_doc_search"])
            self.assertTrue(response.tool_events[0].ok)
            self.assertEqual(response.tool_events[0].payload["query"], "长期不使用")
            self.assertIn("policy search ok", response.assistant_text)

    def test_external_psbc_plugin_zip_can_be_installed_and_enabled(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            plugins_root = root / "plugins_target"
            state_path = root / "plugin_state.json"
            zip_path = self._build_plugin_zip(root)

            manager = PluginManager(plugin_root=plugins_root, state_path=state_path)
            self.assertFalse(any(item["name"] == "psbc_policy" for item in manager.list_plugins()))

            installed = manager.install_plugin(str(zip_path))
            self.assertTrue(installed["ok"])
            self.assertEqual(installed["plugin_name"], "psbc_policy")
            self.assertEqual(installed["source_kind"], "zip")

            plugins = manager.list_plugins()
            self.assertTrue(
                any(
                    item["name"] == "psbc_policy"
                    and not item["enabled"]
                    and item["commercial"]
                    and item["distribution"] == "external"
                    for item in plugins
                )
            )
            self.assertFalse(any(item["name"] == "policy_doc_search" for item in manager.command_specs()))

            enabled = manager.enable_plugin("psbc_policy")
            self.assertTrue(enabled["ok"])
            self.assertTrue(any(item["name"] == "policy_doc_search" for item in manager.command_specs()))

    def test_external_psbc_plugin_executes_without_repo_plugins_imports(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            plugins_root = root / "plugins_target"
            state_path = root / "plugin_state.json"
            zip_path = self._build_plugin_zip(root)
            library_root = root / "policy_runtime"
            docx_path = root / "idle_account_policy.docx"

            document = Document()
            document.add_heading("Idle Account Handling Policy", level=1)
            document.add_paragraph("Long-term unused operations accounts should be locked or restricted.")
            document.save(str(docx_path))

            manager = PluginManager(plugin_root=plugins_root, state_path=state_path)
            installed = manager.install_plugin(str(zip_path))
            self.assertTrue(installed["ok"])

            original_import = builtins.__import__

            def _guarded_import(name, globals=None, locals=None, fromlist=(), level=0):
                if name.startswith("plugins.psbc_policy"):
                    raise AssertionError(f"unexpected repo plugin import: {name}")
                return original_import(name, globals, locals, fromlist, level)

            with patch("builtins.__import__", side_effect=_guarded_import):
                enabled = manager.enable_plugin("psbc_policy")
                self.assertTrue(enabled["ok"])

                imported = manager._tools["policy_doc_import"].handler(
                    path=str(docx_path),
                    library_root=str(library_root),
                )
                self.assertTrue(imported.ok)

                searched = manager._tools["policy_doc_search"].handler(
                    query="unused account lock",
                    library_root=str(library_root),
                    limit=5,
                )
                self.assertTrue(searched.ok)
                self.assertGreaterEqual(int(searched.payload.get("count") or 0), 1)

    def test_external_psbc_plugin_can_be_disabled_and_removed_cleanly(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            plugins_root = root / "plugins_target"
            state_path = root / "plugin_state.json"
            zip_path = self._build_plugin_zip(root)

            manager = PluginManager(plugin_root=plugins_root, state_path=state_path)
            installed = manager.install_plugin(str(zip_path))
            self.assertTrue(installed["ok"])

            enabled = manager.enable_plugin("psbc_policy")
            self.assertTrue(enabled["ok"])
            self.assertTrue(any(item["name"] == "policy_doc_search" for item in manager.command_specs()))
            self.assertTrue(any(item["name"] == "policy_doc_search" for item in manager.tool_specs()))
            self.assertTrue(any(item["name"] == "policy_query" for item in manager.command_specs()))
            self.assertTrue(any(item["name"] == "policy_query" for item in manager.tool_specs()))
            self.assertIn("policy_query", self._openai_model_tool_names(manager))

            disabled = manager.disable_plugin("psbc_policy")
            self.assertTrue(disabled["ok"])
            self.assertFalse(any(item["name"] == "policy_doc_search" for item in manager.command_specs()))
            self.assertFalse(any(item["name"] == "policy_doc_search" for item in manager.tool_specs()))
            self.assertFalse(any(item["name"] == "policy_query" for item in manager.command_specs()))
            self.assertFalse(any(item["name"] == "policy_query" for item in manager.tool_specs()))
            self.assertNotIn("policy_query", self._openai_model_tool_names(manager))

            removed = manager.remove_plugin("psbc_policy")
            self.assertTrue(removed["ok"])
            self.assertFalse((plugins_root / "psbc_policy").exists())
            self.assertFalse(any(item["name"] == "psbc_policy" for item in manager.list_plugins()))
            self.assertNotIn("policy_query", self._openai_model_tool_names(manager))
