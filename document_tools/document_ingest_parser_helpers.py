#!/usr/bin/env python3
"""DOCX, spreadsheet, and PDF parsers for document ingest."""

from __future__ import annotations

import re
import shutil
import tempfile
from pathlib import Path
from typing import Any, Dict, List, Optional, Sequence, Tuple

from document_tools.document_ingest_file_helpers import _truncate


def _extract_docx_title(document: Any, paragraphs: Sequence[str], path: Path) -> str:
    for paragraph in document.paragraphs:
        text = re.sub(r"\s+", " ", paragraph.text or "").strip()
        if not text:
            continue
        style_name = getattr(getattr(paragraph, "style", None), "name", "") or ""
        if "title" in style_name.lower() or "heading" in style_name.lower():
            return text
    for text in paragraphs:
        if 4 <= len(text) <= 80:
            return text
    return path.stem


def parse_docx(path: Path) -> Dict:
    import docx2txt
    from docx import Document

    document = Document(str(path))
    paragraphs = [
        re.sub(r"\s+", " ", paragraph.text or "").strip()
        for paragraph in document.paragraphs
        if (paragraph.text or "").strip()
    ]
    tables = []
    for table_index, table in enumerate(document.tables[:5], start=1):
        rows = []
        for row in table.rows[:10]:
            values = [re.sub(r"\s+", " ", cell.text or "").strip() for cell in row.cells[:10]]
            if any(values):
                rows.append(values)
        if rows:
            tables.append({"table_index": table_index, "rows": rows})
    text_fallback = _truncate(docx2txt.process(str(path)) or "", limit=1200)
    return {
        "title": _extract_docx_title(document, paragraphs, path),
        "paragraph_count": len(paragraphs),
        "paragraph_preview": paragraphs[:12],
        "table_count": len(document.tables),
        "tables": tables,
        "full_text_preview": text_fallback,
    }


def _header_row(values: List[List[str]]) -> Tuple[Optional[int], List[str]]:
    for index, row in enumerate(values[:10], start=1):
        cleaned = [cell for cell in row if cell]
        if len(cleaned) >= 2:
            return index, row
    if values:
        return 1, values[0]
    return None, []


def _normalize_cell_text(value: Any) -> str:
    return "" if value in (None, "") else str(value).strip()


def _column_letter(col: int) -> str:
    col = max(1, col)
    letters = []
    while col:
        col, remainder = divmod(col - 1, 26)
        letters.append(chr(65 + remainder))
    return "".join(reversed(letters))


def _cell_coordinate(row: int, col: int) -> str:
    return f"{_column_letter(col)}{max(1, row)}"


def _collect_openpyxl_merged_cells(sheet) -> List[Dict[str, Any]]:
    merged_cells: List[Dict[str, Any]] = []
    for merged_range in getattr(sheet.merged_cells, "ranges", []) or []:
        min_col, min_row, max_col, max_row = merged_range.bounds
        merged_cells.append(
            {
                "range": str(merged_range),
                "start_cell": _cell_coordinate(min_row, min_col),
                "end_cell": _cell_coordinate(max_row, max_col),
                "start_row": int(min_row),
                "end_row": int(max_row),
                "start_col": int(min_col),
                "end_col": int(max_col),
                "row_span": int(max_row - min_row + 1),
                "col_span": int(max_col - min_col + 1),
                "anchor_value": _normalize_cell_text(sheet.cell(row=min_row, column=min_col).value),
            }
        )
    return merged_cells


def _collect_xlrd_merged_cells(sheet) -> List[Dict[str, Any]]:
    merged_cells: List[Dict[str, Any]] = []
    for row_start0, row_end0, col_start0, col_end0 in getattr(sheet, "merged_cells", []) or []:
        row_start = int(row_start0 + 1)
        row_end = int(row_end0)
        col_start = int(col_start0 + 1)
        col_end = int(col_end0)
        merged_cells.append(
            {
                "range": f"{_cell_coordinate(row_start, col_start)}:{_cell_coordinate(row_end, col_end)}",
                "start_cell": _cell_coordinate(row_start, col_start),
                "end_cell": _cell_coordinate(row_end, col_end),
                "start_row": row_start,
                "end_row": row_end,
                "start_col": col_start,
                "end_col": col_end,
                "row_span": int(row_end - row_start + 1),
                "col_span": int(col_end - col_start + 1),
                "anchor_value": _normalize_cell_text(sheet.cell_value(row_start0, col_start0)),
            }
        )
    return merged_cells


def _sheet_preview(sheet, *, merged_cells: Optional[List[Dict[str, Any]]] = None) -> Dict:
    rows: List[List[str]] = []
    non_empty_rows = 0
    for row in sheet.iter_rows(min_row=1, max_row=min(sheet.max_row or 1, 80), values_only=True):
        normalized = [_normalize_cell_text(cell) for cell in row[:12]]
        if any(normalized):
            non_empty_rows += 1
            rows.append(normalized)
        if len(rows) >= 12:
            break
    header_index, header = _header_row(rows)
    sample_rows = []
    if header_index is not None:
        sample_rows = rows[header_index:header_index + 5]
    key_cells = {}
    for row in range(1, min(sheet.max_row or 1, 6) + 1):
        for col in range(1, min(sheet.max_column or 1, 6) + 1):
            value = sheet.cell(row=row, column=col).value
            if value not in (None, ""):
                key_cells[f"{sheet.cell(row=row, column=col).coordinate}"] = _normalize_cell_text(value)
    merged_cells = merged_cells or []
    return {
        "sheet_name": sheet.title,
        "max_row": int(sheet.max_row or 0),
        "max_column": int(sheet.max_column or 0),
        "non_empty_rows_scanned": non_empty_rows,
        "header_row_index": header_index,
        "header": header,
        "sample_rows": sample_rows,
        "key_cells": key_cells,
        "merged_cell_count": len(merged_cells),
        "merged_cells_preview": merged_cells[:12],
        "merged_cells": merged_cells,
    }


def parse_xlsx(path: Path) -> Dict:
    from openpyxl import load_workbook

    workbook = load_workbook(filename=str(path), read_only=False, data_only=True)
    sheets = []
    for name in workbook.sheetnames[:10]:
        sheet = workbook[name]
        merged_cells = _collect_openpyxl_merged_cells(sheet)
        sheets.append(_sheet_preview(sheet, merged_cells=merged_cells))
    sheet_names = list(workbook.sheetnames)
    workbook.close()
    return {
        "sheet_count": len(sheet_names),
        "sheet_names": sheet_names,
        "sheets": sheets,
    }


def parse_xls(path: Path) -> Dict:
    import xlrd

    cleanup_path: Optional[Path] = None
    try:
        workbook = xlrd.open_workbook(str(path), on_demand=True)
    except OSError:
        safe_dir = Path(tempfile.mkdtemp(prefix="office_xls_"))
        cleanup_path = safe_dir / "source.xls"
        shutil.copyfile(path, cleanup_path)
        workbook = xlrd.open_workbook(str(cleanup_path), on_demand=True)
    sheet_names = workbook.sheet_names()
    sheets = []
    for name in sheet_names[:10]:
        sheet = workbook.sheet_by_name(name)
        rows: List[List[str]] = []
        non_empty_rows = 0
        for row_idx in range(min(sheet.nrows, 80)):
            normalized = []
            for col_idx in range(min(sheet.ncols, 12)):
                value = sheet.cell_value(row_idx, col_idx)
                text = "" if value in (None, "") else str(value).strip()
                normalized.append(text)
            if any(normalized):
                non_empty_rows += 1
                rows.append(normalized)
            if len(rows) >= 12:
                break
        header_index, header = _header_row(rows)
        sample_rows = rows[header_index:header_index + 5] if header_index is not None else []
        key_cells = {}
        for row_idx in range(min(sheet.nrows, 5)):
            for col_idx in range(min(sheet.ncols, 5)):
                value = sheet.cell_value(row_idx, col_idx)
                if value not in (None, ""):
                    col_name = xlrd.formula.colname(col_idx)
                    key_cells[f"{col_name}{row_idx + 1}"] = _normalize_cell_text(value)
        merged_cells = _collect_xlrd_merged_cells(sheet)
        sheets.append(
            {
                "sheet_name": name,
                "max_row": int(sheet.nrows),
                "max_column": int(sheet.ncols),
                "non_empty_rows_scanned": non_empty_rows,
                "header_row_index": header_index,
                "header": header,
                "sample_rows": sample_rows,
                "key_cells": key_cells,
                "merged_cell_count": len(merged_cells),
                "merged_cells_preview": merged_cells[:12],
                "merged_cells": merged_cells,
            }
        )
    workbook.release_resources()
    if cleanup_path is not None:
        shutil.rmtree(cleanup_path.parent, ignore_errors=True)
    return {
        "sheet_count": len(sheet_names),
        "sheet_names": sheet_names,
        "sheets": sheets,
    }


def parse_pdf(path: Path) -> Dict:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    metadata = reader.metadata or {}
    pages = []
    full_text_parts: List[str] = []
    for index, page in enumerate(reader.pages, start=1):
        text = re.sub(r"\s+\n", "\n", page.extract_text() or "")
        text = re.sub(r"\n{3,}", "\n\n", text).strip()
        pages.append(
            {
                "page_number": index,
                "text_preview": _truncate(text, 1200),
            }
        )
        if text:
            full_text_parts.append(text)
    title = str(getattr(metadata, "title", "") or metadata.get("/Title") or path.stem).strip() or path.stem
    return {
        "title": title,
        "page_count": len(reader.pages),
        "pages": pages,
        "full_text_preview": _truncate("\n\n".join(full_text_parts), 4000),
        "metadata": {
            "author": str(getattr(metadata, "author", "") or metadata.get("/Author") or "").strip(),
            "subject": str(getattr(metadata, "subject", "") or metadata.get("/Subject") or "").strip(),
        },
    }
