#!/usr/bin/env python3
"""
附件文档解析。

支持：
- zip/7z 解压
- doc/docx/pdf 解析
- xls/xlsx 解析
"""

from __future__ import annotations

import json
import re
from collections.abc import Sequence
from pathlib import Path
from typing import Any

try:
    import docx2txt
except ModuleNotFoundError:  # pragma: no cover - optional dependency in some environments
    docx2txt = None
from docx import Document

from shared.document_tools.document_ingest_archive import (
    convert_legacy_office,
    extract_7z_safe,
    extract_zip_safe,
)
from shared.document_tools.document_ingest_markdown import render_markdown as _render_markdown_impl
from shared.document_tools.document_ingest_spreadsheets import (
    parse_xls as _parse_xls_impl,
)
from shared.document_tools.document_ingest_spreadsheets import (
    parse_xlsx as _parse_xlsx_impl,
)

SUPPORTED_SUFFIXES = {".zip", ".7z", ".doc", ".docx", ".pdf", ".xls", ".xlsx"}
PARSEABLE_SUFFIXES = {".doc", ".docx", ".pdf", ".xls", ".xlsx"}


def _safe_name(value: str) -> str:
    return re.sub(r"[^0-9A-Za-z\u4e00-\u9fff._-]+", "_", value).strip("_") or "file"


def _truncate(text: str, limit: int = 160) -> str:
    text = re.sub(r"\s+", " ", (text or "")).strip()
    if len(text) <= limit:
        return text
    return text[: limit - 3] + "..."


def _extract_docx_title(document: Document, paragraphs: Sequence[str], path: Path) -> str:
    for paragraph in document.paragraphs:
        text = re.sub(r"\s+", " ", paragraph.text or "").strip()
        if not text:
            continue
        style_name = getattr(getattr(paragraph, "style", None), "name", "") or ""
        if "title" in style_name.lower() or "heading" in style_name.lower():
            return text
    for text in paragraphs:
        if 4 <= len(text) <= 80:
            return text
    return path.stem


def parse_docx(path: Path) -> dict:
    document = Document(str(path))
    paragraphs = [
        re.sub(r"\s+", " ", paragraph.text or "").strip()
        for paragraph in document.paragraphs
        if (paragraph.text or "").strip()
    ]
    tables = []
    for table_index, table in enumerate(document.tables[:5], start=1):
        rows = []
        for row in table.rows[:10]:
            values = [re.sub(r"\s+", " ", cell.text or "").strip() for cell in row.cells[:10]]
            if any(values):
                rows.append(values)
        if rows:
            tables.append({"table_index": table_index, "rows": rows})
    if docx2txt is not None:
        text_fallback = _truncate(docx2txt.process(str(path)) or "", limit=1200)
    else:
        fallback_chunks = list(paragraphs)
        for table in tables:
            for row in table.get("rows") or []:
                fallback_chunks.append(" | ".join(cell for cell in row if cell))
        text_fallback = _truncate("\n".join(fallback_chunks), limit=1200)
    return {
        "title": _extract_docx_title(document, paragraphs, path),
        "paragraph_count": len(paragraphs),
        "paragraph_preview": paragraphs[:12],
        "table_count": len(document.tables),
        "tables": tables,
        "full_text_preview": text_fallback,
    }


def parse_xlsx(path: Path) -> dict:
    return _parse_xlsx_impl(path)


def parse_xls(path: Path) -> dict:
    return _parse_xls_impl(path)


def parse_pdf(path: Path) -> dict:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    metadata = reader.metadata or {}
    pages = []
    full_text_parts: list[str] = []
    for index, page in enumerate(reader.pages, start=1):
        text = re.sub(r"\s+\n", "\n", page.extract_text() or "")
        text = re.sub(r"\n{3,}", "\n\n", text).strip()
        pages.append(
            {
                "page_number": index,
                "text_preview": _truncate(text, 1200),
            }
        )
        if text:
            full_text_parts.append(text)
    title = (
        str(getattr(metadata, "title", "") or metadata.get("/Title") or path.stem).strip()
        or path.stem
    )
    return {
        "title": title,
        "page_count": len(reader.pages),
        "pages": pages,
        "full_text_preview": _truncate("\n\n".join(full_text_parts), 4000),
        "metadata": {
            "author": str(getattr(metadata, "author", "") or metadata.get("/Author") or "").strip(),
            "subject": str(
                getattr(metadata, "subject", "") or metadata.get("/Subject") or ""
            ).strip(),
        },
    }


def render_markdown(file_type: str, content: dict, *, title: str) -> str:
    return _render_markdown_impl(file_type=file_type, content=content, title=title)


def _write_markdown_artifact(markdown_root: Path, source_name: str, markdown_text: str) -> str:
    markdown_root.mkdir(parents=True, exist_ok=True)
    output_path = markdown_root / f"{_safe_name(Path(source_name).stem)}.md"
    output_path.write_text(markdown_text, encoding="utf-8")
    return str(output_path)


def _build_merge_metadata(file_type: str, content: dict, *, source_name: str) -> dict[str, Any]:
    return {
        "source_name": source_name,
        "file_type": file_type,
        "sheet_count": int(content.get("sheet_count") or 0),
        "sheet_names": content.get("sheet_names") or [],
        "sheets": [
            {
                "sheet_name": sheet.get("sheet_name"),
                "max_row": int(sheet.get("max_row") or 0),
                "max_column": int(sheet.get("max_column") or 0),
                "merged_cell_count": int(sheet.get("merged_cell_count") or 0),
                "merged_cells": sheet.get("merged_cells") or [],
            }
            for sheet in content.get("sheets", []) or []
        ],
    }


def _write_json_artifact(
    output_root: Path, source_name: str, payload: dict[str, Any], *, suffix: str
) -> str:
    output_root.mkdir(parents=True, exist_ok=True)
    output_path = output_root / f"{_safe_name(Path(source_name).stem)}{suffix}"
    output_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    return str(output_path)


def _flatten_parseable_files(
    paths: Sequence[Path],
    extract_root: Path,
    convert_root: Path,
    errors: list[dict] | None = None,
) -> list[Path]:
    resolved: list[Path] = []
    for path in paths:
        suffix = path.suffix.lower()
        if suffix == ".zip":
            extracted = extract_zip_safe(path, extract_root)
            resolved.extend(
                _flatten_parseable_files(extracted, extract_root, convert_root, errors=errors)
            )
            continue
        if suffix == ".7z":
            try:
                extracted = extract_7z_safe(path, extract_root)
            except Exception as exc:
                if errors is not None:
                    errors.append({"path": str(path), "stage": "extract", "error": str(exc)})
                continue
            resolved.extend(
                _flatten_parseable_files(extracted, extract_root, convert_root, errors=errors)
            )
            continue
        if suffix not in PARSEABLE_SUFFIXES:
            continue
        if suffix == ".doc":
            try:
                resolved.append(convert_legacy_office(path, convert_root))
            except Exception as exc:
                if errors is not None:
                    errors.append({"path": str(path), "stage": "convert", "error": str(exc)})
            continue
        resolved.append(path)
    return resolved


def summarize_document_content(file_type: str, content: dict) -> str:
    if file_type == "docx":
        title = content.get("title") or ""
        paragraph_preview = content.get("paragraph_preview") or []
        lead = paragraph_preview[0] if paragraph_preview else ""
        return _truncate("；".join(part for part in [title, lead] if part), 220)
    if file_type == "pdf":
        title = content.get("title") or ""
        lead = content.get("full_text_preview") or ""
        return _truncate("；".join(part for part in [title, lead] if part), 220)
    if file_type in {"xlsx", "xls"}:
        sheet_names = content.get("sheet_names") or []
        parts = []
        if sheet_names:
            parts.append(f"工作表: {', '.join(sheet_names[:4])}")
        first_sheet = (content.get("sheets") or [{}])[0]
        header = [cell for cell in first_sheet.get("header", []) if cell]
        if header:
            parts.append(f"表头: {', '.join(header[:6])}")
        return _truncate("；".join(parts), 220)
    return ""


def ingest_document(path: Path, runtime_root: Path | None = None) -> dict:
    path = Path(path)
    runtime_root = runtime_root or (path.parent / ".document_ingest")
    extract_root = runtime_root / "extracted"
    convert_root = runtime_root / "converted"
    markdown_root = runtime_root / "markdown"
    merge_metadata_root = runtime_root / "merge_metadata"
    extract_root.mkdir(parents=True, exist_ok=True)
    convert_root.mkdir(parents=True, exist_ok=True)
    markdown_root.mkdir(parents=True, exist_ok=True)
    merge_metadata_root.mkdir(parents=True, exist_ok=True)

    errors: list[dict] = []
    parse_targets = _flatten_parseable_files(
        [path],
        extract_root=extract_root,
        convert_root=convert_root,
        errors=errors,
    )
    documents = []
    for target in parse_targets:
        suffix = target.suffix.lower()
        try:
            if suffix == ".docx":
                file_type = "docx"
                content = parse_docx(target)
                title = str(content.get("title") or target.stem)
            elif suffix == ".xlsx":
                file_type = "xlsx"
                content = parse_xlsx(target)
                title = target.stem
            elif suffix == ".xls":
                file_type = "xls"
                content = parse_xls(target)
                title = target.stem
            elif suffix == ".pdf":
                file_type = "pdf"
                content = parse_pdf(target)
                title = str(content.get("title") or target.stem)
            else:
                continue
            markdown_text = render_markdown(file_type, content, title=title)
            markdown_path = _write_markdown_artifact(markdown_root, target.name, markdown_text)
            merge_metadata_path = None
            if file_type in {"xlsx", "xls"}:
                merge_metadata_path = _write_json_artifact(
                    merge_metadata_root,
                    target.name,
                    _build_merge_metadata(file_type, content, source_name=target.name),
                    suffix=".merge_metadata.json",
                )
            documents.append(
                {
                    "path": str(target),
                    "name": target.name,
                    "file_type": file_type,
                    "content": content,
                    "content_summary": summarize_document_content(file_type, content),
                    "markdown_path": markdown_path,
                    "markdown_preview": _truncate(markdown_text, 2000),
                    "merge_metadata_path": merge_metadata_path,
                }
            )
        except Exception as exc:
            errors.append({"path": str(target), "stage": "parse", "error": str(exc)})

    return {
        "source_path": str(path),
        "source_name": path.name,
        "source_suffix": path.suffix.lower(),
        "resolved_document_count": len(documents),
        "documents": documents,
        "errors": errors,
    }


def ingest_documents(paths: Sequence[Path], runtime_root: Path) -> list[dict]:
    result = []
    for path in paths:
        result.append(ingest_document(Path(path), runtime_root=runtime_root))
    return result


def dumps_json(data: dict) -> str:
    return json.dumps(data, ensure_ascii=False, indent=2)
