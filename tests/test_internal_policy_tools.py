import sys
import tempfile
import unittest
import json
from pathlib import Path
from unittest.mock import patch

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from shared.document_tools.internal_policy_tools import InternalPolicyTools
from internal_policy_docs.library import list_policy_documents, register_policy_markdown_corpus
from shared.document_tools.policy_query import policy_query_compact_queries, policy_query_terms

class InternalPolicyToolsTest(unittest.TestCase):
    def test_default_library_auto_mounts_multiple_corpora_with_metadata(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            library_root = root / "policy_runtime"
            racs_root = root / "RACS" / "data"

            (racs_root / "psbc" / "markdown").mkdir(parents=True, exist_ok=True)
            (racs_root / "psbc" / "markdown" / "中国邮政储蓄银行数据安全管理办法.md").write_text(
                "# 中国邮政储蓄银行数据安全管理办法\n\n邮储内部数据安全管理要求。\n",
                encoding="utf-8",
            )
            (racs_root / "laws").mkdir(parents=True, exist_ok=True)
            (racs_root / "laws" / "中华人民共和国数据安全法.md").write_text(
                "# 中华人民共和国数据安全法\n\n国家层面的数据安全法律依据。\n",
                encoding="utf-8",
            )
            (racs_root / "gb" / "markdown").mkdir(parents=True, exist_ok=True)
            (racs_root / "gb" / "markdown" / "GB_T22239-2019.md").write_text(
                "# GB/T 22239-2019\n\n网络安全等级保护标准。\n",
                encoding="utf-8",
            )

            with patch("internal_policy_docs.library.DEFAULT_LIBRARY_ROOT", library_root), patch(
                "internal_policy_docs.library.DEFAULT_RACS_DATA_ROOT", racs_root
            ):
                listed = list_policy_documents(library_root=str(library_root), limit=20)

            self.assertTrue(listed["ok"])
            corpora = listed["corpora"]
            corpus_names = [item["corpus_name"] for item in corpora]
            self.assertEqual(corpus_names, ["psbc", "laws", "gb"])
            self.assertEqual(corpora[0]["authority_level"], "bank_internal_policy")
            self.assertEqual(corpora[1]["authority_level"], "national_law")
            self.assertEqual(corpora[2]["authority_level"], "national_standard")
            self.assertEqual(corpora[1]["authority_rank"], 100)
            self.assertEqual(corpora[0]["scope"], "psbc_internal")
            self.assertEqual(corpora[1]["domain_tags"][0], "laws")

    def test_import_list_search_and_read(self):
        tools = InternalPolicyTools()
        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            library_root = root / "policy_runtime"
            docx_path = root / "travel_policy.docx"

            document = Document()
            document.add_heading("Travel Reimbursement Policy", level=1)
            document.add_paragraph("Employees must submit travel reimbursement within five working days.")
            document.add_paragraph("Hotel reimbursement follows the internal grade standard.")
            document.save(str(docx_path))

            imported = tools.policy_doc_import(str(docx_path), library_root=str(library_root))
            self.assertTrue(imported["ok"])
            self.assertEqual(imported["imported_count"], 1)
            imported_doc = imported["documents"][0]
            self.assertTrue(imported_doc["doc_id"])
            self.assertTrue(imported_doc["markdown_path"])

            listed = tools.policy_doc_list(library_root=str(library_root))
            self.assertTrue(listed["ok"])
            self.assertEqual(listed["count"], 1)
            self.assertEqual(listed["documents"][0]["doc_id"], imported_doc["doc_id"])

            searched = tools.policy_doc_search("reimbursement hotel", library_root=str(library_root))
            self.assertTrue(searched["ok"])
            self.assertEqual(searched["count"], 1)
            self.assertEqual(searched["documents"][0]["doc_id"], imported_doc["doc_id"])

            read_back = tools.policy_doc_read(doc_id=imported_doc["doc_id"], library_root=str(library_root))
            self.assertTrue(read_back["ok"])
            self.assertIn("Travel Reimbursement Policy", read_back["text"])
            self.assertIn("five working days", read_back["text"])

    def test_external_markdown_corpus_list_search_and_read(self):
        tools = InternalPolicyTools()
        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            library_root = root / "policy_runtime"
            corpus_root = root / "psbc" / "markdown"
            corpus_root.mkdir(parents=True, exist_ok=True)
            markdown_path = corpus_root / "运营数据中心差旅费管理细则（测试版）.md"
            markdown_path.write_text(
                "# 运营数据中心差旅费管理细则\n\n出差报销应在返程后五个工作日内完成提交。\n",
                encoding="utf-8",
            )

            registered = register_policy_markdown_corpus(
                corpus_root,
                library_root=str(library_root),
                corpus_name="psbc_markdown_test",
            )
            self.assertTrue(registered["ok"])

            listed = tools.policy_doc_list(library_root=str(library_root))
            self.assertTrue(listed["ok"])
            self.assertEqual(listed["count"], 1)
            self.assertEqual(listed["documents"][0]["source_kind"], "external_markdown")

            searched = tools.policy_doc_search("差旅费 报销", library_root=str(library_root))
            self.assertTrue(searched["ok"])
            self.assertEqual(searched["count"], 1)
            found_doc = searched["documents"][0]
            self.assertEqual(found_doc["corpus_name"], "psbc_markdown_test")

            read_back = tools.policy_doc_read(doc_id=found_doc["doc_id"], library_root=str(library_root))
            self.assertTrue(read_back["ok"])
            self.assertIn("差旅费管理细则", read_back["text"])
            self.assertIn("五个工作日", read_back["text"])

    def test_external_markdown_corpus_saved_as_wsl_unc_still_works_on_linux(self):
        tools = InternalPolicyTools()
        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            library_root = root / "policy_runtime"
            corpus_root = root / "psbc" / "markdown"
            corpus_root.mkdir(parents=True, exist_ok=True)
            markdown_path = corpus_root / "中国邮政储蓄银行数据安全管理办法（测试版）.md"
            markdown_path.write_text(
                "# 中国邮政储蓄银行数据安全管理办法（测试版）\n\n数据分类分级应遵循统一管理要求。\n",
                encoding="utf-8",
            )

            registered = register_policy_markdown_corpus(
                corpus_root,
                library_root=str(library_root),
                corpus_name="psbc_markdown_test",
            )
            self.assertTrue(registered["ok"])

            index_path = library_root / "library_index.json"
            payload = json.loads(index_path.read_text(encoding="utf-8"))
            payload["corpora"][0]["markdown_root"] = rf"\\wsl.localhost\Ubuntu{corpus_root.as_posix().replace('/', '\\')}"
            payload["corpora"][0]["source_root"] = rf"\\wsl.localhost\Ubuntu{corpus_root.parent.as_posix().replace('/', '\\')}"
            index_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")

            listed = tools.policy_doc_list(library_root=str(library_root))
            self.assertTrue(listed["ok"])
            self.assertEqual(listed["count"], 1)

            searched = tools.policy_doc_search("数据分类分级", library_root=str(library_root))
            self.assertTrue(searched["ok"])
            self.assertEqual(searched["count"], 1)

            found_doc = searched["documents"][0]
            read_back = tools.policy_doc_read(doc_id=found_doc["doc_id"], library_root=str(library_root))
            self.assertTrue(read_back["ok"])
            self.assertIn("统一管理要求", read_back["text"])

    def test_legacy_psbc_markdown_corpus_name_is_normalized_to_psbc(self):
        tools = InternalPolicyTools()
        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            library_root = root / "policy_runtime"
            corpus_root = root / "psbc" / "markdown"
            corpus_root.mkdir(parents=True, exist_ok=True)
            markdown_path = corpus_root / "中国邮政储蓄银行数据安全管理办法（测试版）.md"
            markdown_path.write_text(
                "# 中国邮政储蓄银行数据安全管理办法（测试版）\n\n邮储内部数据安全管理要求。\n",
                encoding="utf-8",
            )

            registered = register_policy_markdown_corpus(
                corpus_root,
                library_root=str(library_root),
                corpus_name="psbc_markdown",
            )
            self.assertTrue(registered["ok"])

            listed = tools.policy_doc_list(library_root=str(library_root))
            self.assertTrue(listed["ok"])
            self.assertEqual(listed["corpora"][0]["corpus_name"], "psbc")

            searched = tools.policy_doc_search("邮储 数据安全", library_root=str(library_root), limit=3)
            self.assertTrue(searched["ok"])
            self.assertGreaterEqual(searched["count"], 1)
            self.assertEqual(searched["documents"][0]["corpus_name"], "psbc")

    def test_search_prefers_normative_policy_over_training_and_audit_noise(self):
        tools = InternalPolicyTools()
        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            library_root = root / "policy_runtime"
            corpus_root = root / "psbc" / "markdown"
            corpus_root.mkdir(parents=True, exist_ok=True)

            (corpus_root / "中国邮政储蓄银行信息系统用户账号和权限管理实施细则（2024年修订版）.md").write_text(
                "# 中国邮政储蓄银行信息系统用户账号和权限管理实施细则（2024年修订版）\n"
                "最小授权原则。动态控制原则。长期未登录账号最长为一年应锁定或限制访问。\n",
                encoding="utf-8",
            )
            (corpus_root / "6.中国邮政储蓄银行运维安全堡垒系统管理规程.md").write_text(
                "# 6.中国邮政储蓄银行运维安全堡垒系统管理规程\n"
                "宣贯培训。工作中确需访问生产系统才可申请账号。长期未登录最长90天自动锁定。\n",
                encoding="utf-8",
            )
            (corpus_root / "2025年信用卡核心系统审计底稿-权限问题反馈.md").write_text(
                "# 2025年信用卡核心系统审计底稿-权限问题反馈\n"
                "审计发现存在权限管理薄弱和长期闲置账号问题。\n",
                encoding="utf-8",
            )

            registered = register_policy_markdown_corpus(
                corpus_root,
                library_root=str(library_root),
                corpus_name="psbc_markdown_test",
            )
            self.assertTrue(registered["ok"])

            searched = tools.policy_doc_search("权限管理 长期闲置", library_root=str(library_root), limit=5)
            self.assertTrue(searched["ok"])
            self.assertGreaterEqual(searched["count"], 2)

            first = searched["documents"][0]
            self.assertEqual(first["doc_kind"], "governance_policy")
            self.assertEqual(first["doc_group"], "governance_base")
            self.assertIn("信息系统用户账号和权限管理实施细则", first["title"])

            doc_kinds = [item["doc_kind"] for item in searched["documents"]]
            self.assertIn("audit_workpaper", doc_kinds)
            self.assertTrue(doc_kinds.index("governance_policy") < doc_kinds.index("audit_workpaper"))

    def test_search_returns_query_aware_evidence_summary_and_term_hits(self):
        tools = InternalPolicyTools()
        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            library_root = root / "policy_runtime"
            corpus_root = root / "psbc" / "markdown"
            corpus_root.mkdir(parents=True, exist_ok=True)

            (
                corpus_root / "中国邮政储蓄银行信息系统用户账号和权限管理实施细则（2024年修订版）.md"
            ).write_text(
                "# 中国邮政储蓄银行信息系统用户账号和权限管理实施细则（2024年修订版）\n"
                "最小授权原则。动态控制原则。长期未登录账号应锁定或限制访问。\n",
                encoding="utf-8",
            )
            (
                corpus_root / "中国邮政储蓄银行运维安全堡垒系统管理规程（2024年修订版）.md"
            ).write_text(
                "# 中国邮政储蓄银行运维安全堡垒系统管理规程（2024年修订版）\n"
                "确需访问生产系统才可申请账号。不再需要时应申请注销或调整。\n",
                encoding="utf-8",
            )

            registered = register_policy_markdown_corpus(
                corpus_root,
                library_root=str(library_root),
                corpus_name="psbc_markdown_test",
            )
            self.assertTrue(registered["ok"])

            searched = tools.policy_doc_search("核心业务应用运维管控系统:长期闲置账号制度依据", library_root=str(library_root), limit=5)
            self.assertTrue(searched["ok"])
            self.assertGreaterEqual(searched["count"], 1)

            first = searched["documents"][0]
            self.assertGreaterEqual(first["query_term_hits"], 1)
            self.assertTrue(first["matched_terms"])
            self.assertTrue(first["evidence_summary"])
            self.assertLessEqual(len(first["evidence_summary"]), 180)
            self.assertIn(first["matched_terms"][0], first["evidence_summary"])
            self.assertTrue(first["source_queries"])
            self.assertGreaterEqual(first["query_variant_hits"], 1)
            self.assertTrue(searched["query_variants"])

    def test_search_fuses_query_variants_for_long_policy_questions(self):
        tools = InternalPolicyTools()
        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            library_root = root / "policy_runtime"
            corpus_root = root / "psbc" / "markdown"
            corpus_root.mkdir(parents=True, exist_ok=True)

            (
                corpus_root / "中国邮政储蓄银行信息系统用户账号和权限管理实施细则（2024年修订版）.md"
            ).write_text(
                "# 中国邮政储蓄银行信息系统用户账号和权限管理实施细则（2024年修订版）\n"
                "动态控制原则。人员岗位发生变化时应及时调整其账号和权限，避免出现多余账号、多余权限或权责不对应。\n"
                "长期闲置账号、长期未登录账号最长为一年应锁定或限制访问，并收回其权限。\n",
                encoding="utf-8",
            )
            (
                corpus_root / "17-中国邮政储蓄银行客户信息平台-调研表.md"
            ).write_text(
                "# 17-中国邮政储蓄银行客户信息平台-调研表\n"
                "调研确认信息，包含账号数量、联系人、系统背景等描述。\n",
                encoding="utf-8",
            )

            registered = register_policy_markdown_corpus(
                corpus_root,
                library_root=str(library_root),
                corpus_name="psbc_markdown_test",
            )
            self.assertTrue(registered["ok"])

            searched = tools.policy_doc_search(
                "邮政储蓄银行核心业务应用运维管控系统:存在30名用户并非按需申请，长期闲置问题，制度依据是什么？",
                library_root=str(library_root),
                limit=5,
            )
            self.assertTrue(searched["ok"])
            self.assertGreaterEqual(searched["count"], 1)

            first = searched["documents"][0]
            self.assertIn("信息系统用户账号和权限管理实施细则", first["title"])
            self.assertGreaterEqual(first["query_variant_hits"], 2)
            self.assertGreaterEqual(len(searched["query_variants"]), 2)

    def test_policy_query_helpers_compact_audit_finding_to_short_queries(self):
        finding = "经查，你中心部分外包服务提供商尽职调查执行不到位，未见外包服务提供商信息安全管理能力、财务状况、资质能力等尽职调查记录，请说明制度依据。"

        terms = policy_query_terms(finding, limit=12)
        queries = policy_query_compact_queries(finding, limit=4)

        self.assertIn("外包服务提供商", terms)
        self.assertIn("尽职调查", terms)
        self.assertIn("外包服务提供商 尽职调查", queries)
        self.assertTrue(all("经查" not in query for query in queries))
        self.assertTrue(all("请说明" not in query for query in queries))

    def test_policy_query_helpers_keep_scenario_and_drop_answer_request_suffix(self):
        finding = "针对外包服务提供商尽职调查不到位的问题，请给出制度依据、问题定性和责任环节。"

        terms = policy_query_terms(finding, limit=12)
        queries = policy_query_compact_queries(finding, limit=4)

        self.assertIn("外包服务提供商", terms)
        self.assertIn("尽职调查", terms)
        self.assertIn("外包服务提供商 尽职调查", queries)
        self.assertTrue(all("问题定性和责任环节" not in query for query in queries))

    def test_policy_query_helpers_keep_surface_subject_for_hidden_topics(self):
        least_privilege = "针对外包人员权限与职责不匹配的问题，请给出制度依据、问题定性和责任环节。"
        reporting = "针对外包活动和驻场外包人员未按季度报送的问题，请给出制度依据、问题定性和责任环节。"

        least_privilege_queries = policy_query_compact_queries(least_privilege, limit=4)
        reporting_queries = policy_query_compact_queries(reporting, limit=4)

        self.assertTrue(any("外包人员" in query and "职责" in query for query in least_privilege_queries))
        self.assertTrue(any("外包活动" in query and "驻场外包人员" in query for query in reporting_queries))

    def test_search_prefers_compact_audit_queries_over_irrelevant_law_titles(self):
        tools = InternalPolicyTools()
        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            library_root = root / "policy_runtime"

            psbc_root = root / "psbc" / "markdown"
            laws_root = root / "laws"
            psbc_root.mkdir(parents=True, exist_ok=True)
            laws_root.mkdir(parents=True, exist_ok=True)

            (psbc_root / "中国邮政储蓄银行信息科技外包管理办法（2024年修订版）.md").write_text(
                "# 中国邮政储蓄银行信息科技外包管理办法（2024年修订版）\n"
                "在拟使用外包服务提供商的信息科技外包服务前，应按照外包服务提供商尽职调查表开展尽职调查。\n"
                "调查内容应包括财务情况、业务能力、风险管理与内控情况、业务连续性管理情况等，并保留相关调查资料。\n",
                encoding="utf-8",
            )
            (laws_root / "国家金库条例.md").write_text(
                "# 国家金库条例\n"
                "国家金库管理应当依法开展，不涉及外包服务提供商尽职调查要求。\n",
                encoding="utf-8",
            )

            register_policy_markdown_corpus(psbc_root, library_root=str(library_root))
            register_policy_markdown_corpus(laws_root, library_root=str(library_root))

            searched = tools.policy_doc_search(
                "经查，你中心部分外包服务提供商尽职调查执行不到位，未见外包服务提供商信息安全管理能力、财务状况、资质能力等尽职调查记录，请说明制度依据。",
                library_root=str(library_root),
                limit=5,
            )
            self.assertTrue(searched["ok"])
            self.assertGreaterEqual(searched["count"], 1)
            self.assertIn("外包服务提供商 尽职调查", searched["query_variants"])
            self.assertIn("信息科技外包管理办法", searched["documents"][0]["title"])

    def test_search_exposes_corpus_metadata_and_prefers_upper_level_corpus_for_basis_query(self):
        tools = InternalPolicyTools()
        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            library_root = root / "policy_runtime"

            psbc_root = root / "psbc" / "markdown"
            laws_root = root / "laws"
            regulatory_root = root / "regulatory" / "raw"
            psbc_root.mkdir(parents=True, exist_ok=True)
            laws_root.mkdir(parents=True, exist_ok=True)
            regulatory_root.mkdir(parents=True, exist_ok=True)

            (psbc_root / "中国邮政储蓄银行数据安全管理办法.md").write_text(
                "# 中国邮政储蓄银行数据安全管理办法\n\n本办法规定本行内部数据安全管理职责和要求。\n",
                encoding="utf-8",
            )
            (laws_root / "中华人民共和国数据安全法.md").write_text(
                "# 中华人民共和国数据安全法\n\n国家数据安全工作的基础法律依据。\n",
                encoding="utf-8",
            )
            (regulatory_root / "银行保险机构数据安全管理办法.md").write_text(
                "# 银行保险机构数据安全管理办法\n\n银行保险机构应建立数据安全管理体系和制度机制。\n",
                encoding="utf-8",
            )

            register_policy_markdown_corpus(psbc_root, library_root=str(library_root))
            register_policy_markdown_corpus(laws_root, library_root=str(library_root))
            register_policy_markdown_corpus(regulatory_root, library_root=str(library_root))

            searched = tools.policy_doc_search("数据安全管理办法 上位依据", library_root=str(library_root), limit=5)
            self.assertTrue(searched["ok"])
            self.assertGreaterEqual(searched["count"], 2)

            first = searched["documents"][0]
            self.assertIn(first["corpus_authority_level"], {"national_law", "financial_regulation"})
            self.assertIn(first["corpus_name"], {"laws", "regulatory"})
            self.assertIn("corpus_metadata", first)
            self.assertIn("authority_level", first["corpus_metadata"])
            self.assertGreater(int(first["corpus_query_boost"]), 0)

    def test_search_prefers_psbc_for_internal_policy_query(self):
        tools = InternalPolicyTools()
        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            library_root = root / "policy_runtime"

            psbc_root = root / "psbc" / "markdown"
            laws_root = root / "laws"
            regulatory_root = root / "regulatory" / "raw"
            psbc_root.mkdir(parents=True, exist_ok=True)
            laws_root.mkdir(parents=True, exist_ok=True)
            regulatory_root.mkdir(parents=True, exist_ok=True)

            (psbc_root / "中国邮政储蓄银行数据安全管理办法.md").write_text(
                "# 中国邮政储蓄银行数据安全管理办法\n\n邮储银行应按照本行制度开展数据分类分级和权限控制。\n",
                encoding="utf-8",
            )
            (laws_root / "中华人民共和国数据安全法.md").write_text(
                "# 中华人民共和国数据安全法\n\n国家建立数据安全管理制度。\n",
                encoding="utf-8",
            )
            (regulatory_root / "银行保险机构数据安全管理办法.md").write_text(
                "# 银行保险机构数据安全管理办法\n\n银行保险机构应建立数据安全管理机制。\n",
                encoding="utf-8",
            )

            register_policy_markdown_corpus(psbc_root, library_root=str(library_root))
            register_policy_markdown_corpus(laws_root, library_root=str(library_root))
            register_policy_markdown_corpus(regulatory_root, library_root=str(library_root))

            searched = tools.policy_doc_search("邮储 数据安全 管理要求", library_root=str(library_root), limit=5)
            self.assertTrue(searched["ok"])
            self.assertGreaterEqual(searched["count"], 1)

            first = searched["documents"][0]
            self.assertEqual(first["corpus_authority_level"], "bank_internal_policy")
            self.assertEqual(first["corpus_name"], "psbc")
            self.assertGreaterEqual(int(first["corpus_query_boost"]), 30)

    def test_search_treats_postal_savings_bank_full_name_as_internal_intent(self):
        tools = InternalPolicyTools()
        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            library_root = root / "policy_runtime"

            psbc_root = root / "psbc" / "markdown"
            laws_root = root / "laws"
            psbc_root.mkdir(parents=True, exist_ok=True)
            laws_root.mkdir(parents=True, exist_ok=True)

            (psbc_root / "中国邮政储蓄银行数据安全管理办法.md").write_text(
                "# 中国邮政储蓄银行数据安全管理办法\n\n本办法规定本行内部数据安全管理职责和要求。\n",
                encoding="utf-8",
            )
            (laws_root / "银行保险机构数据安全管理办法.md").write_text(
                "# 银行保险机构数据安全管理办法\n\n银行保险机构应建立数据安全管理机制。\n",
                encoding="utf-8",
            )

            register_policy_markdown_corpus(psbc_root, library_root=str(library_root))
            register_policy_markdown_corpus(laws_root, library_root=str(library_root))

            searched = tools.policy_doc_search("中国邮政储蓄银行数据安全管理办法", library_root=str(library_root), limit=5)
            self.assertTrue(searched["ok"])
            self.assertGreaterEqual(searched["count"], 1)

            first = searched["documents"][0]
            self.assertEqual(first["corpus_name"], "psbc")
            self.assertEqual(first["corpus_authority_level"], "bank_internal_policy")

    def test_search_keeps_upper_basis_visible_but_still_hits_psbc_for_neutral_policy_title(self):
        tools = InternalPolicyTools()
        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            library_root = root / "policy_runtime"

            psbc_root = root / "psbc" / "markdown"
            laws_root = root / "laws"
            regulatory_root = root / "regulatory" / "raw"
            psbc_root.mkdir(parents=True, exist_ok=True)
            laws_root.mkdir(parents=True, exist_ok=True)
            regulatory_root.mkdir(parents=True, exist_ok=True)

            (psbc_root / "中国邮政储蓄银行数据安全管理办法.md").write_text(
                "# 中国邮政储蓄银行数据安全管理办法\n\n本办法规定本行内部数据安全管理职责和要求。\n",
                encoding="utf-8",
            )
            (laws_root / "中华人民共和国数据安全法.md").write_text(
                "# 中华人民共和国数据安全法\n\n国家数据安全工作的基础法律依据。\n",
                encoding="utf-8",
            )
            (regulatory_root / "银行保险机构数据安全管理办法.md").write_text(
                "# 银行保险机构数据安全管理办法\n\n银行保险机构应建立数据安全管理体系和制度机制。\n",
                encoding="utf-8",
            )

            register_policy_markdown_corpus(psbc_root, library_root=str(library_root))
            register_policy_markdown_corpus(laws_root, library_root=str(library_root))
            register_policy_markdown_corpus(regulatory_root, library_root=str(library_root))

            searched = tools.policy_doc_search("数据安全管理办法", library_root=str(library_root), limit=5)
            self.assertTrue(searched["ok"])
            self.assertGreaterEqual(searched["count"], 3)

            top_three = searched["documents"][:3]
            self.assertTrue(any(item["corpus_name"] == "psbc" for item in top_three))
            self.assertTrue(any(item["corpus_name"] in {"laws", "regulatory"} for item in top_three))

    def test_search_prefers_internal_operational_policy_for_domain_audit_query(self):
        tools = InternalPolicyTools()
        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            library_root = root / "policy_runtime"

            psbc_root = root / "psbc" / "markdown"
            laws_root = root / "laws"
            psbc_root.mkdir(parents=True, exist_ok=True)
            laws_root.mkdir(parents=True, exist_ok=True)

            (psbc_root / "中国邮政储蓄银行运维安全堡垒系统管理规程（2024年修订版）.md").write_text(
                "# 中国邮政储蓄银行运维安全堡垒系统管理规程（2024年修订版）\n"
                "责任部门负责审核辖内人员申请运维安全堡垒系统的用户账号和权限是否符合最小授权原则。\n"
                "根据工作职责授予用户账号完成工作所需最小权限，避免出现权责不一致和过度授权。\n",
                encoding="utf-8",
            )
            (laws_root / "市场主体登记管理条例实施细则.md").write_text(
                "# 市场主体登记管理条例实施细则\n"
                "办理登记事项应当遵循诚实守信原则，依法履行授权程序。\n",
                encoding="utf-8",
            )

            register_policy_markdown_corpus(psbc_root, library_root=str(library_root))
            register_policy_markdown_corpus(laws_root, library_root=str(library_root))

            searched = tools.policy_doc_search("最小授权原则 外包人员 权限", library_root=str(library_root), limit=5)
            self.assertTrue(searched["ok"])
            self.assertGreaterEqual(searched["count"], 1)

            first = searched["documents"][0]
            self.assertEqual(first["corpus_name"], "psbc")
            self.assertIn("运维安全堡垒系统管理规程", first["title"])
            self.assertGreaterEqual(int(first["corpus_query_boost"]), 30)

    def test_search_exact_internal_policy_title_ranks_exact_document_first(self):
        tools = InternalPolicyTools()
        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            library_root = root / "policy_runtime"

            psbc_root = root / "psbc" / "markdown"
            laws_root = root / "laws"
            psbc_root.mkdir(parents=True, exist_ok=True)
            laws_root.mkdir(parents=True, exist_ok=True)

            (psbc_root / "中国邮政储蓄银行运维安全堡垒系统管理规程（2024年修订版）.md").write_text(
                "# 中国邮政储蓄银行运维安全堡垒系统管理规程（2024年修订版）\n"
                "UKey与用户账号一对一绑定，责任部门负责规范使用和定期检查。\n",
                encoding="utf-8",
            )
            (laws_root / "危险废物经营许可证管理办法.md").write_text(
                "# 危险废物经营许可证管理办法\n"
                "经营单位应当规范许可证使用和授权管理。\n",
                encoding="utf-8",
            )

            register_policy_markdown_corpus(psbc_root, library_root=str(library_root))
            register_policy_markdown_corpus(laws_root, library_root=str(library_root))

            searched = tools.policy_doc_search("中国邮政储蓄银行运维安全堡垒系统管理规程", library_root=str(library_root), limit=5)
            self.assertTrue(searched["ok"])
            self.assertGreaterEqual(searched["count"], 1)

            first = searched["documents"][0]
            self.assertEqual(first["corpus_name"], "psbc")
            self.assertIn("运维安全堡垒系统管理规程", first["title"])
