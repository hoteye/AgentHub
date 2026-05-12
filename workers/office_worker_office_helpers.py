"""Office document operation helpers for the Office worker."""

from __future__ import annotations

import re
import shutil
import tempfile
from pathlib import Path
from typing import Any, Dict, List, Optional


def _truncate(text: str, limit: int = 220) -> str:
    text = re.sub(r"\s+", " ", (text or "")).strip()
    if len(text) <= limit:
        return text
    return text[: limit - 3] + "..."


def _read_xls_summary(
    path: str,
    *,
    max_rows: int = 12,
    max_cols: int = 12,
) -> Dict[str, Any]:
    import xlrd

    cleanup_path: Optional[Path] = None
    try:
        workbook = xlrd.open_workbook(path, on_demand=True)
    except OSError:
        safe_dir = Path(tempfile.mkdtemp(prefix="office_xls_"))
        cleanup_path = safe_dir / "source.xls"
        shutil.copyfile(path, cleanup_path)
        workbook = xlrd.open_workbook(str(cleanup_path), on_demand=True)
    sheets: List[Dict[str, Any]] = []
    for name in workbook.sheet_names()[:10]:
        sheet = workbook.sheet_by_name(name)
        rows: List[List[str]] = []
        non_empty_rows = 0
        max_scan_rows = min(sheet.nrows, max_rows)
        max_scan_cols = min(sheet.ncols, max_cols)
        for row_idx in range(max_scan_rows):
            normalized = []
            for col_idx in range(max_scan_cols):
                value = sheet.cell_value(row_idx, col_idx)
                text = "" if value in (None, "") else str(value).strip()
                normalized.append(text)
            if any(normalized):
                non_empty_rows += 1
                rows.append(normalized)
        header = rows[0] if rows else []
        sample_rows = rows[1:min(len(rows), 6)] if len(rows) > 1 else []
        key_cells = {}
        for row_idx in range(min(sheet.nrows, 5)):
            for col_idx in range(min(sheet.ncols, 5)):
                value = sheet.cell_value(row_idx, col_idx)
                if value not in (None, ""):
                    col_name = xlrd.formula.colname(col_idx)
                    key_cells[f"{col_name}{row_idx + 1}"] = str(value).strip()
        sheets.append(
            {
                "sheet_name": name,
                "max_row": int(sheet.nrows),
                "max_column": int(sheet.ncols),
                "non_empty_rows_scanned": non_empty_rows,
                "header": header,
                "sample_rows": sample_rows,
                "key_cells": key_cells,
            }
        )
    workbook.release_resources()
    if cleanup_path is not None:
        shutil.rmtree(cleanup_path.parent, ignore_errors=True)
    return {
        "ok": True,
        "skill": "read_xls_summary",
        "path": str(Path(path).resolve()),
        "sheet_count": len(sheets),
        "sheet_names": [sheet["sheet_name"] for sheet in sheets],
        "sheets": sheets,
    }


def _read_docx_outline(
    path: str,
    *,
    max_paragraphs: int = 20,
    max_tables: int = 5,
) -> Dict[str, Any]:
    from docx import Document

    document = Document(path)
    paragraphs = [
        " ".join((paragraph.text or "").split())
        for paragraph in document.paragraphs
        if (paragraph.text or "").strip()
    ]
    tables: List[Dict[str, Any]] = []
    for table_index, table in enumerate(document.tables[:max_tables], start=1):
        rows: List[List[str]] = []
        for row in table.rows[:10]:
            values = [" ".join((cell.text or "").split()) for cell in row.cells[:10]]
            if any(values):
                rows.append(values)
        if rows:
            tables.append({"table_index": table_index, "rows": rows})

    title = Path(path).stem
    for paragraph in document.paragraphs:
        text = " ".join((paragraph.text or "").split())
        if not text:
            continue
        style_name = getattr(getattr(paragraph, "style", None), "name", "") or ""
        if "title" in style_name.lower() or "heading" in style_name.lower():
            title = text
            break
    else:
        if paragraphs:
            title = paragraphs[0]

    return {
        "ok": True,
        "skill": "read_docx_outline",
        "path": str(Path(path).resolve()),
        "title": title,
        "paragraph_count": len(paragraphs),
        "paragraph_preview": paragraphs[:max_paragraphs],
        "table_count": len(document.tables),
        "tables": tables,
    }


def _read_xlsx_summary(
    path: str,
    *,
    max_rows: int = 12,
    max_cols: int = 12,
) -> Dict[str, Any]:
    from openpyxl import load_workbook

    workbook = load_workbook(filename=path, read_only=True, data_only=True)
    sheets: List[Dict[str, Any]] = []
    for name in workbook.sheetnames[:10]:
        sheet = workbook[name]
        rows: List[List[str]] = []
        non_empty_rows = 0
        for row in sheet.iter_rows(
            min_row=1,
            max_row=min(sheet.max_row or 1, max_rows),
            values_only=True,
        ):
            normalized = ["" if cell is None else str(cell).strip() for cell in row[:max_cols]]
            if any(normalized):
                non_empty_rows += 1
                rows.append(normalized)
        header = rows[0] if rows else []
        sample_rows = rows[1: min(len(rows), 6)] if len(rows) > 1 else []
        key_cells = {}
        for row_idx in range(1, min(sheet.max_row or 1, 6) + 1):
            for col_idx in range(1, min(sheet.max_column or 1, 6) + 1):
                value = sheet.cell(row=row_idx, column=col_idx).value
                if value not in (None, ""):
                    key_cells[sheet.cell(row=row_idx, column=col_idx).coordinate] = str(value).strip()
        sheets.append(
            {
                "sheet_name": name,
                "max_row": int(sheet.max_row or 0),
                "max_column": int(sheet.max_column or 0),
                "non_empty_rows_scanned": non_empty_rows,
                "header": header,
                "sample_rows": sample_rows,
                "key_cells": key_cells,
            }
        )
    workbook.close()
    return {
        "ok": True,
        "skill": "read_xlsx_summary",
        "path": str(Path(path).resolve()),
        "sheet_count": len(sheets),
        "sheet_names": [sheet["sheet_name"] for sheet in sheets],
        "sheets": sheets,
    }


def _summarize_docx_outline(result: Dict[str, Any]) -> str:
    title = result.get("title") or ""
    paragraphs = result.get("paragraph_preview") or []
    lead = paragraphs[0] if paragraphs else ""
    return _truncate("；".join(part for part in [title, lead] if part))


def _summarize_xlsx_summary(result: Dict[str, Any]) -> str:
    parts: List[str] = []
    if result.get("sheet_names"):
        parts.append(f"工作表: {', '.join(result['sheet_names'][:4])}")
    first_sheet = (result.get("sheets") or [{}])[0]
    header = [cell for cell in first_sheet.get("header", []) if cell]
    if header:
        parts.append(f"表头: {', '.join(header[:6])}")
    return _truncate("；".join(parts))


def _read_xlsx_sheet(
    path: str,
    *,
    sheet_name: Optional[str] = None,
    max_rows: int = 20,
    max_cols: int = 12,
) -> Dict[str, Any]:
    from openpyxl import load_workbook

    workbook = load_workbook(filename=path, read_only=True, data_only=True)
    target_name = sheet_name or workbook.sheetnames[0]
    sheet = workbook[target_name]
    rows: List[List[str]] = []
    for row in sheet.iter_rows(
        min_row=1,
        max_row=min(sheet.max_row or 1, max_rows),
        values_only=True,
    ):
        rows.append(["" if cell is None else str(cell).strip() for cell in row[:max_cols]])
    workbook.close()
    return {
        "ok": True,
        "skill": "read_xlsx_sheet",
        "path": str(Path(path).resolve()),
        "sheet_name": target_name,
        "rows": rows,
        "row_count": len(rows),
        "max_row": int(sheet.max_row or 0),
        "max_column": int(sheet.max_column or 0),
    }


def _update_xlsx_cells(
    path: str,
    *,
    sheet_name: str,
    updates: List[Dict[str, Any]],
    save_as: Optional[str] = None,
) -> Dict[str, Any]:
    from openpyxl import load_workbook

    workbook = load_workbook(filename=path)
    sheet = workbook[sheet_name]
    applied: List[Dict[str, Any]] = []
    for update in updates:
        cell = str(update["cell"]).strip()
        value = update.get("value")
        sheet[cell] = value
        applied.append({"cell": cell, "value": value})

    output_path = Path(save_as) if save_as else Path(path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    workbook.save(str(output_path))
    workbook.close()
    return {
        "ok": True,
        "skill": "update_xlsx_cells",
        "path": str(output_path.resolve()),
        "sheet_name": sheet_name,
        "updates": applied,
        "updated_count": len(applied),
    }


def _create_xlsx(
    path: str,
    *,
    sheets: Optional[List[Dict[str, Any]]] = None,
) -> Dict[str, Any]:
    from openpyxl import Workbook

    workbook = Workbook()
    default_sheet = workbook.active
    default_sheet.title = (sheets[0]["name"] if sheets else "Sheet1")
    if sheets:
        for index, spec in enumerate(sheets):
            if index == 0:
                sheet = default_sheet
            else:
                sheet = workbook.create_sheet(title=spec["name"])
            for row in spec.get("rows", []):
                sheet.append(list(row))
    output_path = Path(path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    workbook.save(str(output_path))
    workbook.close()
    return {
        "ok": True,
        "skill": "create_xlsx",
        "path": str(output_path.resolve()),
        "sheet_names": [sheet.title for sheet in workbook.worksheets] if False else [spec["name"] for spec in sheets] if sheets else ["Sheet1"],
    }


def _create_docx(
    path: str,
    *,
    title: Optional[str] = None,
    paragraphs: Optional[List[str]] = None,
) -> Dict[str, Any]:
    from docx import Document

    document = Document()
    if title:
        document.add_heading(title, level=0)
    for paragraph in paragraphs or []:
        document.add_paragraph(paragraph)
    output_path = Path(path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    return {
        "ok": True,
        "skill": "create_docx",
        "path": str(output_path.resolve()),
        "title": title,
        "paragraph_count": len(paragraphs or []),
    }


def _append_docx_paragraphs(
    path: str,
    *,
    paragraphs: List[str],
    heading: Optional[str] = None,
    save_as: Optional[str] = None,
) -> Dict[str, Any]:
    from docx import Document

    document = Document(path)
    if heading:
        document.add_heading(heading, level=1)
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    output_path = Path(save_as) if save_as else Path(path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    return {
        "ok": True,
        "skill": "append_docx_paragraphs",
        "path": str(output_path.resolve()),
        "heading": heading,
        "paragraph_count": len(paragraphs),
    }


def _inspect_office_file(path: str) -> Dict[str, Any]:
    target = Path(path)
    suffix = target.suffix.lower()
    file_kind = "unknown"
    if suffix in {".xlsx", ".xls"}:
        file_kind = "spreadsheet"
    elif suffix in {".docx", ".doc", ".pdf"}:
        file_kind = "document"
    return {
        "ok": True,
        "skill": "inspect_office_file",
        "path": str(target.resolve()),
        "exists": target.exists(),
        "name": target.name,
        "suffix": suffix,
        "file_kind": file_kind,
        "size_bytes": int(target.stat().st_size) if target.exists() else None,
    }
